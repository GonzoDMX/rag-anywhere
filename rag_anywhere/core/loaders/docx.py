# rag_anywhere/core/loaders/docx.py

from pathlib import Path
from typing import Dict, Any

from .base import DocumentLoader


class DocxLoader(DocumentLoader):
    """Loader for Microsoft Word documents"""
    
    def __init__(self):
        try:
            import docx
            self.docx = docx
        except ImportError:
            raise ImportError(
                "DOCX support requires 'python-docx' package. "
                "Install with: pip install python-docx"
            )
    
    def load(self, file_path: Path) -> str:
        """Load DOCX file and extract text"""
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        try:
            doc = self.docx.Document(str(file_path))
            
            # Extract text from paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Extract text from tables
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text for cell in row.cells)
                    if row_text.strip():
                        table_texts.append(row_text)
            
            # Combine paragraphs and tables
            all_text = paragraphs
            if table_texts:
                all_text.extend(['', '--- Tables ---', ''] + table_texts)
            
            return "\n\n".join(all_text)
        except Exception as e:
            raise ValueError(f"Error loading DOCX {file_path}: {e}")
    
    def supports(self, file_path: Path) -> bool:
        """Check if file is DOCX"""
        return file_path.suffix.lower() in ['.docx', '.doc']
    
    def get_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from DOCX"""
        stat = file_path.stat()
        metadata = {
            'filename': file_path.name,
            'file_size': stat.st_size,
            'file_type': file_path.suffix.lower(),
            'mime_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        }
        
        try:
            doc = self.docx.Document(str(file_path))
            core_props = doc.core_properties
            
            if core_props.title:
                metadata['title'] = core_props.title
            if core_props.author:
                metadata['author'] = core_props.author
            if core_props.subject:
                metadata['subject'] = core_props.subject
            if core_props.created:
                metadata['created'] = core_props.created.isoformat()
            if core_props.modified:
                metadata['modified'] = core_props.modified.isoformat()
            
            # Count paragraphs
            metadata['num_paragraphs'] = len([p for p in doc.paragraphs if p.text.strip()])
            metadata['num_tables'] = len(doc.tables)
        except Exception:
            pass  # If metadata extraction fails, just return basic info
        
        return metadata
